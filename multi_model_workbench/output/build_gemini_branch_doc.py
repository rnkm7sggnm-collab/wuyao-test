from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT_PATH = Path("/Users/heer/Documents/🍑 NO1/multi_model_workbench/output/雾窑_Gemini支线扩写资料.docx")


def set_font(run, size=11, bold=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc, text="", style=None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    if text:
        run = paragraph.add_run(text)
        set_font(run)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        set_font(run)


def add_numbers(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        set_font(run)


def add_heading(doc, text, level):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        set_font(run, size=16 if level == 1 else 13 if level == 2 else 12, bold=True)
    return paragraph


def set_cell_text(cell, text, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_font(run, size=10, bold=bold)


def build_doc():
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    for style_name in ["Normal", "List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)

    title = doc.add_heading("《雾窑》支线扩写资料包", level=0)
    for run in title.runs:
        set_font(run, size=20, bold=True)

    add_heading(doc, "1. 项目一句话简介", 1)
    add_para(doc, "《雾窑》是一款本土化封闭空间悬疑、多结局互动游戏：故事发生在暴雨封村的青瓷庄，玩家围绕二十年前窑火旧案展开证物推理。")
    add_para(doc, "主线围绕失散之子陆景川的复仇展开，玩家需要在封闭环境中查清真相、识破误导，并决定复仇是否继续。")

    add_heading(doc, "2. 故事核心", 1)
    add_para(doc, "《雾窑》的真正核心不是单纯命案，而是二十年前青瓷庄为了利益、名声和沉默，把何青山的死亡伪装成窑火事故。二十年后，被偷走身份的孩子陆景川回到青瓷庄，用复仇逼所有人重新面对真相。")

    add_heading(doc, "3. 已有主线设定", 1)
    add_bullets(doc, [
        "地点：南方山区半荒废古村“青瓷庄”。",
        "封闭条件：暴雨、塌方、吊桥冲毁、断电断网。",
        "表面事件：文旅公司准备改造非遗民宿。",
        "真实矛盾：二十年前“窑火事故”其实是杀人毁证。",
        "玩家目标：活下来，查清旧案，阻止复仇继续。",
    ])

    add_heading(doc, "4. 主要人物", 1)
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["人物", "公开身份", "隐藏秘密", "剧情功能", "可扩写支线方向"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True)

    rows = [
        ["林知夏", "玩家，外来编剧。", "父亲留下旧采访笔记，和旧案调查有关。", "玩家视角、调查入口、证物整合者。", "父亲调查线、笔记缺页、外来者与本地沉默的冲突。"],
        ["陆怀民", "青瓷庄老村长。", "掩盖旧案并抱走何青山的孩子。", "旧案核心知情者，临死遗言误导和提示玩家。", "村庄名声、地方权力、抱走孩子的真实心理。"],
        ["周启明", "文旅负责人。", "挪用公款并借火毁账。", "第一案死者，揭开钱和账本问题。", "文旅开发、利益洗白、账本和汇款单线索。"],
        ["陈伯远", "退休法医。", "伪造尸检结论。", "第二案死者，证明何青山不是烧死。", "忏悔录音、旧尸检报告、临死前想坦白。"],
        ["沈月娥", "青瓷庄最后一位女瓷匠。", "知道丈夫不是意外死亡，可能知道陆景川身世，但没有杀人。", "最大嫌疑人，实际旧案幸存者。", "误解、沉默、保护儿子、洗清嫌疑。"],
        ["陆景川", "民宿项目本地负责人。", "何青山之子，当前真凶和复仇者。", "可靠协助者到最终真凶的反转。", "身份暴露、复仇动机、能否被玩家阻止。"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_text(cell, text)

    add_heading(doc, "5. 二十年前旧案", 1)
    add_numbers(doc, [
        "何青山发现周启明挪用厂款，也发现陆怀民私卖老窑址利益。",
        "举报前夜，争执导致何青山死亡。",
        "涉案者点燃窑房，制造“窑火失控”的假象。",
        "陈伯远把尸检写成意外烧亡。",
        "周启明借火毁账。",
        "陆怀民抱走何青山的儿子，让他以陆景川身份长大。",
    ])

    add_heading(doc, "6. 当前案件", 1)
    add_heading(doc, "第一案：账房中毒", 2)
    add_para(doc, "周启明死亡，釉料粉末指向沈月娥，实际是陆景川栽赃。此案用于制造第一层误导，并引出账本、汇款单和文旅利益线。")
    add_heading(doc, "第二案：尸检报告", 2)
    add_para(doc, "陈伯远窒息，报告被改成“生前钝器伤”，旧案不再是火灾。此案让玩家确认何青山起火前已经死亡。")
    add_heading(doc, "第三案：祠堂遗言", 2)
    add_para(doc, "陆怀民重伤，临死说“不是她，是孩子”，真正指向陆景川。此案连接旧照片、出生证明和陆景川真实身份。")

    add_heading(doc, "7. 核心推理链", 1)
    add_numbers(doc, [
        "窑火不是天灾。",
        "何青山起火前已死。",
        "三名关键人物共同掩盖。",
        "沈月娥被栽赃。",
        "陆景川是何青山之子。",
        "当前命案是复仇。",
    ])

    add_heading(doc, "8. 已有六章流程", 1)
    chapters = [
        ("1. 抵达青瓷庄", "章节目标是让玩家认识人物、进入封闭环境。玩家主要行动是探索老宅、观察青瓷杯、确认暴雨塌方和断电断网。推动的信息是：这不是普通文旅晚宴，而是一场被精心安排的重聚。"),
        ("2. 账房尸体", "章节目标是调查第一案并学习证物系统。玩家主要行动是检查旧账房、找到账本、汇款单和釉料粉末。推动的信息是：周启明与旧案中的钱和账本有关，沈月娥可能被栽赃。"),
        ("3. 尸检报告", "章节目标是证明二十年前的死因有问题。玩家主要行动是调查偏房、寻找陈伯远录音和旧尸检报告。推动的信息是：何青山不是单纯烧死，旧案存在人为掩盖。"),
        ("4. 祠堂遗言", "章节目标是破解陆怀民遗言的误导。玩家主要行动是分析“不是她，是孩子”，调查祠堂照片和陆怀民藏物。推动的信息是：真正的“孩子”可能不是许曼青，而是陆景川。"),
        ("5. 废窑暗格", "章节目标是找到旧案核心证据。玩家主要行动是探索废窑、找到举报信、出生证明、旧照片和采访笔记缺页。推动的信息是：陆景川是何青山之子，当前命案与复仇直接相连。"),
        ("6. 最后对质", "章节目标是还原旧案并决定结局。玩家主要行动是出示证物、指认真凶、选择公开真相或默许复仇。推动的信息是：陆景川是当前凶手，但他也是旧案制造出的受害者。"),
    ]
    for title_text, body in chapters:
        add_heading(doc, title_text, 2)
        add_para(doc, body)

    add_heading(doc, "9. 已有结局", 1)
    add_bullets(doc, [
        "A 真相公开：阻止陆景川，旧案重查，废窑成为纪念空间。",
        "B 替罪羊：沈月娥被误认凶手，陆景川离开，复仇未停。",
        "C 复仇继续：玩家知道真凶但证据不足，真相再次被埋。",
        "D 坏结局：连续推理失败，玩家被困在废窑。",
    ])

    add_heading(doc, "10. 希望 Gemini 帮忙扩写的方向", 1)
    add_para(doc, "请基于以上设定，帮助扩写《雾窑》的支线剧情。要求不要推翻主线真相，不新增过多主要角色，不改变陆景川是复仇者、何青山旧案是核心真相的设定。支线应该服务于证物推理、人物动机、玩家选择和多结局分支。")
    add_para(doc, "请 Gemini 优先扩写以下支线：")
    add_numbers(doc, [
        "沈月娥支线：从最大嫌疑人逐渐变成旧案幸存者，重点是误解、沉默、保护与洗清嫌疑。",
        "陆景川支线：从可靠协助者逐渐暴露为复仇者，重点是身份、复仇动机、是否能被阻止。",
        "林知夏父亲支线：通过旧采访笔记、录音或照片，揭开上一代外来调查者未能公开的真相。",
        "村庄沉默支线：扩写青瓷庄为什么集体沉默，重点是利益链、村庄名声、非遗保护和地方权力。",
        "证物支线：围绕六只青瓷杯、烧焦账本、尸检报告、旧照片、出生证明、举报信、采访笔记等证物设计更多调查节点。",
        "结局支线：扩写不同证物完整度和玩家选择如何影响 A/B/C/D 四个结局。",
    ])

    add_heading(doc, "11. 给 Gemini 的输出要求", 1)
    add_para(doc, "希望 Gemini 输出时遵守：")
    add_bullets(doc, [
        "不要写成长篇小说正文。",
        "优先输出支线结构。",
        "每条支线包括：触发条件、相关人物、关键场景、关键证物、玩家选择、可能后果。",
        "保持本土化悬疑风格。",
        "保持封闭空间压迫感。",
        "不要模仿《无人生还》的童谣杀人结构。",
        "不要新增复杂密室和过多死亡。",
        "所有支线都要回到“二十年前窑火旧案”这个核心。",
    ])

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT_PATH)
