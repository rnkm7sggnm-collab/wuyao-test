from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
MD_PATH = BASE_DIR / "final_multi_model_synthesis.md"
DOCX_PATH = BASE_DIR / "final_multi_model_synthesis.docx"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    style_tokens = {
        "Heading 1": (16, "2E74B5", 18, 10),
        "Heading 2": (13, "2E74B5", 14, 7),
        "Heading 3": (12, "1F4D78", 10, 5),
    }
    for name, (size, color, before, after) in style_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("《雾窑》多模型回答融合整理稿")
    set_run_font(run, size=9, color="666666")


def add_title(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=0, after=12, line=1.25)
    run = paragraph.add_run(text)
    set_run_font(run, size=20, bold=True, color="1F4D78")


def add_horizontal_rule(doc):
    paragraph = doc.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D8DEE8")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    set_paragraph_spacing(paragraph, before=4, after=10, line=1.0)


def add_paragraph_with_bold(doc, text, style=None):
    paragraph = doc.add_paragraph(style=style)
    set_paragraph_spacing(paragraph, after=4 if style in ("List Bullet", "List Number") else 6)
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)
    return paragraph


def build_docx():
    doc = Document()
    style_document(doc)

    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    first_title_done = False
    pending = []

    def flush_pending():
        if pending:
            text = " ".join(part.strip() for part in pending if part.strip())
            if text:
                add_paragraph_with_bold(doc, text)
            pending.clear()

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            flush_pending()
            continue

        if stripped == "---":
            flush_pending()
            add_horizontal_rule(doc)
            continue

        if stripped.startswith("# "):
            flush_pending()
            text = stripped[2:].strip()
            if not first_title_done:
                add_title(doc, text)
                first_title_done = True
            else:
                doc.add_paragraph(text, style="Heading 1")
            continue

        if stripped.startswith("## "):
            flush_pending()
            doc.add_paragraph(stripped[3:].strip(), style="Heading 1")
            continue

        if stripped.startswith("### "):
            flush_pending()
            doc.add_paragraph(stripped[4:].strip(), style="Heading 2")
            continue

        if stripped.startswith("- "):
            flush_pending()
            add_paragraph_with_bold(doc, stripped[2:].strip(), style="List Bullet")
            continue

        ordered_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if ordered_match:
            flush_pending()
            add_paragraph_with_bold(doc, ordered_match.group(1), style="List Number")
            continue

        pending.append(stripped)

    flush_pending()
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_docx()
    print(DOCX_PATH)
